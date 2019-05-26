import time
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
from datetime import datetime, timedelta
import pandas as pd
import os
from docx import Document
import requests
from retry import retry

class CrawlerCB:

    headers = {
        "Host": "www.chinabond.com.cn",
        "User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/70.0.3538.67 Safari/537.36",
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,image/webp,image/apng,*/*;q=0.8",
        "Referer": "http://www.chinabond.com.cn/",
        "Accept-Encoding": "gzip, deflate",
        "Accept-Language": "zh-CN,zh;q=0.9,en;q=0.8",
    }

    db_path = "C:/Users/l_cry/Desktop/docs/"
    #db_path = "ftp://192.168.87.103/docs/"
    def __init__(self):
        options = Options()
        # options.add_argument('-headless')
        self.browser = webdriver.Chrome(options=options)
        self.browser.implicitly_wait(30)
        self.timewait = 5

    # 判断文件是否需要下载
    @staticmethod
    def isdownload(type_name, file_title):
        if type_name in ("企业债"):
            if file_title.find('评级报告') !=-1 or (file_title.find("募集说明书") != -1 and file_title.find("募集说明书摘要") == -1):
                return True
            else: return False
        elif type_name in ("政策性银行债"):
            if file_title.find('主协议') != -1 or file_title.find("要素一览") != -1 or file_title.find('发行办法') !=-1 or \
                    file_title.find('发行文件') !=-1 or file_title.find("函") != -1 :
                return True
            else: return False
        elif type_name in ("资产支持证券"):
            if file_title.find('发行说明') != -1 or file_title.find("发行公告") != -1 or file_title.find('发行办法') != -1 or \
                    file_title.find("承销团成员") != -1 or file_title.find('评级报告') != -1:
                return True
            else: return False
        elif type_name in ("地方政府债"):
            if file_title.find('信息披露文件') != -1 or file_title.find("有关事项") != -1 or file_title.find('评级报告') != -1:
                return True
            else: return False
        else:
            return  False

    @retry(tries=3, delay=200)
    def __parse_bond(self, bond, type_name, Date, sub_name=None):
        if sub_name is None:
            dir_path = self.db_path+type_name
        else:
            dir_path = self.db_path+type_name+"/"+sub_name

        timestr = bond.find_element_by_xpath("span").text
        timestr = timestr.split()[0].replace('/', '-')
        timetime = datetime.strptime(timestr, "%Y-%m-%d")  # 当前公告发布日期时间
        content = bond.find_element_by_xpath("span/a").get_attribute("title")
        content = content.replace('（', '(').replace('）', ')')
        file_notes = {}
        if timetime == Date:
            res = requests.get(bond.find_element_by_xpath("span/a").get_attribute("href"), headers=self.headers)
            if res.status_code != 200:
                raise Exception("http error")
            if 'text/html' in res.headers['Content-Type']:
                bond.find_element_by_xpath("span/a").click()
                time.sleep(self.timewait)
                self.browser.switch_to.window(self.browser.window_handles[-1])
                files = self.browser.find_elements_by_xpath("//ul[@class='zw_malei3']/li")
                for file in files:
                    file_title = file.find_element_by_xpath("span/a").get_attribute("title")
                    if self.isdownload(type_name, file_title):
                        filt_time = file.find_element_by_xpath("p").text
                        r = requests.get(file.find_element_by_xpath("span/a").get_attribute("href"), headers=self.headers)
                        file_path = dir_path+"/"+timestr+"/"+content
                        if not os.path.exists(file_path):
                            os.makedirs(file_path)

                        with open(file_path+"/"+file_title, "wb") as code:
                            code.write(r.content)

                        file_notes[content] = file_path
                self.browser.close()
                self.browser.switch_to.window(self.browser.window_handles[0])
            else:
                return 0
                # file_path = dir_path + "/" + timestr+ "/" +content
                # if not os.path.exists(file_path):
                #     os.makedirs(file_path)
                #
                # with open(file_path + "/" + content+".doc", "wb") as code:
                #     code.write(res.content)
        return file_notes

    @retry(tries=3, delay=200)
    def __parse_state_debt(self, bond, type_name, Date, sub_name=None):
        if sub_name is None:
            dir_path = self.db_path+type_name
        else:
            dir_path = self.db_path+type_name+"/"+sub_name
        timestr = bond.find_element_by_xpath("span").text
        timestr = timestr.split()[0].replace('/', '-')
        timetime = datetime.strptime(timestr, "%Y-%m-%d")
        content = bond.find_element_by_xpath("span/a").get_attribute("title")
        content = content.replace('（', '(').replace('）', ')')
        if timetime == Date:
            document = Document()
            bond.find_element_by_xpath("span/a").click()
            time.sleep(self.timewait)
            self.browser.switch_to.window(self.browser.window_handles[-1])
            article_title = self.browser.find_element_by_xpath("//div[@class='zw_main2']/strong").text
            document.add_heading(article_title, 3)
            article_subtitle = self.browser.find_element_by_xpath("//div[@class='zw_main2']/font").text
            document.add_paragraph(article_subtitle)
            article_content = self.browser.find_element_by_xpath("//div[@id='hiddenContent']").text
            document.add_paragraph(article_content)
            article_time = self.browser.find_element_by_xpath("//div[@class='zw_main6']").text
            document.add_paragraph(article_time)
            self.browser.close()
            self.browser.switch_to.window(self.browser.window_handles[0])
            file_path = dir_path+"/"+timestr
            if not os.path.exists(file_path):
                os.makedirs(file_path)
            document.save(file_path+"/"+article_title+".docx")
        return 0

    ########################################################################
    # 根据日期抓取信息
    def crawbyDate(self, Date):
        if type(Date) is str:
            Date = pd.to_datetime(Date)
        self.browser.get("http://www.chinabond.com.cn/Channel/21000")
        dl_list = self.browser.find_elements_by_xpath("//div[@class='fx_subnav']/dl")
        for dl in dl_list:
            type_name = dl.find_element_by_xpath("dt").text.strip()
            print(type_name)
            if type_name == "全部":
                pass
            elif type_name == "国债":
                dl.find_element_by_xpath("dt").click()
                if dl.find_element_by_xpath("dd").get_attribute("style") == "display: none;":
                    dl.find_element_by_xpath("dt").click()
                a_list = dl.find_elements_by_xpath("dd/a")
                for a in a_list:
                    a.click()
                    time.sleep(self.timewait)
                    sub_name = a.text
                    self.browser.switch_to.frame("ffrIframe")
                    bonds = self.browser.find_elements_by_xpath("//li[@class='liqxd1']")
                    for bond in bonds:
                        self.__parse_state_debt(bond, type_name, Date, sub_name)

                    self.browser.switch_to.default_content()

            elif type_name in ("地方政府债", "企业债","非银行金融债"):
                dl.click()
                time.sleep(self.timewait)
                self.browser.switch_to.frame("ffrIframe")
                bonds = self.browser.find_elements_by_xpath("//li[@class='liqxd1']")
                for bond in bonds:
                    self.__parse_bond(bond, type_name, Date)
                self.browser.switch_to.default_content()

            elif type_name in ("政策性银行债", "资产支持证券"):
                dl.find_element_by_xpath("dt").click()
                if  dl.find_element_by_xpath("dd").get_attribute("style") == "display: none;":
                    dl.find_element_by_xpath("dt").click()
                a_list = dl.find_elements_by_xpath("dd/a")
                for a in a_list:
                    time.sleep(3)
                    a.click()
                    time.sleep(self.timewait)
                    sub_name = a.get_attribute("text")
                    self.browser.switch_to.frame("ffrIframe")
                    bonds = self.browser.find_elements_by_xpath("//li[@class='liqxd1']")
                    if bonds[0].text.strip() != "暂无相关文件":
                        for bond in bonds:
                            self.__parse_bond(bond, type_name, Date, sub_name)
                    self.browser.switch_to.default_content()
            else:
                pass
        self.browser.close()


if __name__ == '__main__':
    crawlercb = CrawlerCB()
    now = datetime.now()
    yesterday = now - timedelta(1)
    tomorrow = now + timedelta(1)
    try:
        crawlercb.crawbyDate(now.strftime("%Y-%m-%d"))
    except:
        print("中债登爬取失败")




